import json
import logging
import os
import re
import tempfile
import time
import traceback
import unicodedata
import urllib.error
import urllib.request
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from openai import OpenAI
from openpyxl import Workbook
import fitz
from pypdf import PdfReader
from sqlalchemy.orm import Session
from typhoon_ocr import prepare_ocr_messages

from .audit import record_audit
from .config import get_settings
from .models import AuditEvent, DictionaryTerm, Job, JobStatus, SpellcheckFinding
from .storage import download_file, upload_file

MAX_CHARS = 15000
PAGE_CONCURRENCY = 3  # ลดจาก 10 เพื่อไม่ให้ยิง OpenRouter พร้อมกันจนชน rate limit (free tier 20 req/min)
JOB_TIMEOUT_SECONDS = 600  # ห้ามงานค้างสถานะ PROCESSING เกิน 10 นาที
logger = logging.getLogger(__name__)

THAI_MARKS = frozenset("\u0e31\u0e34\u0e35\u0e36\u0e37\u0e38\u0e39\u0e3a\u0e47\u0e48\u0e49\u0e4a\u0e4b\u0e4c\u0e4d\u0e4e")
THAI_BASES = frozenset(chr(codepoint) for codepoint in range(0x0E01, 0x0E2F))

SYSTEM_PROMPT = """You are a careful Thai proofreading assistant. Return concise Markdown only.

กติกา:
- เป้าหมายคือคำสะกดผิด คำพิมพ์ตก และคำซ้ำเท่านั้น
- รายงานเฉพาะคำที่น่าจะสะกดผิดหรือ OCR อ่านตัวอักษรผิดจริง
- อย่ารายงานการปรับสำนวน ความกระชับ หรือคำที่ถูกอยู่แล้ว
- อย่ารายงานปัญหาเว้นวรรค เช่น "สิ่ง แวดล้อม" -> "สิ่งแวดล้อม"
- อย่ารายงานการเปลี่ยนตัวพิมพ์ใหญ่/เล็กของคำอังกฤษ
- อย่ารายงานการเติมหรือลบเครื่องหมายจุด (.) หรือเครื่องหมายจุลภาค (,) เนื่องจากภาษาไทยไม่ใช้จุดปิดท้ายประโยคหรือเครื่องหมายจุลภาคคั่นประโยค
- อย่าแก้ชื่อเฉพาะ ชื่อหลักสูตร ชื่อหน่วยงาน ตัวย่อ รหัส เลข พ.ศ. หรือคำอังกฤษ ถ้าไม่มั่นใจ
- ห้ามใส่รายการที่คำที่พบเหมือนคำที่แนะนำ
- ถ้าไม่พบ ให้ตอบว่า: ไม่พบคำผิดชัดเจน
- รูปแบบแต่ละรายการ: - `คำที่พบ` -> `คำที่แนะนำ` : เหตุผลสั้นๆ
- เอกสารนี้คือเล่มหลักสูตรมหาวิทยาลัย ซึ่งจัดทำล่วงหน้าก่อนเปิดใช้งานจริง ปีการศึกษาในเอกสารอาจเป็นปีในอนาคตได้ เช่น 2569, 2570"""


def normalize_thai(text: str) -> str:
    """Normalize Thai text to fix decomposed characters from PDF extraction.

    1. Apply NFC normalization for general Unicode composition.
    2. Manually compose สระอำ: U+0E4D (นิคหิต ํ) + U+0E32 (สระอา า) → U+0E33 (สระอำ ำ).
       This pair is NOT a canonical decomposition in Unicode, so NFC alone won't fix it.
    """
    text = unicodedata.normalize("NFC", text)
    # ponytail: PDF text layers sometimes insert spaces before zero-width Thai marks.
    text = re.sub(r"(?<=[\u0e01-\u0e3a\u0e40-\u0e4e]) +(?=[\u0e31\u0e34-\u0e3a\u0e47-\u0e4e])", "", text)
    # ponytail: PDF text extraction loses glyph width when a zero-width nikhahit is mapped to a space.
    text = re.sub(r"(?<=[\u0e01-\u0e2e\u0e31\u0e34-\u0e3a\u0e47-\u0e4e]) \u0e32", "\u0e33", text)
    # Fix spaces before สระอำ or decomposed สระอำ
    text = text.replace(" \u0e33", "\u0e33")
    text = text.replace(" \u0e4d\u0e32", "\u0e33")
    text = text.replace("\u0e4d\u0e32", "\u0e33")  # ํา → ำ
    return text


def has_unicode_artifacts(text: str) -> bool:
    """Detect broken PDF/OCR text without relying on a word dictionary.

    Custom PDF fonts commonly leak private-use glyphs, invisible formatting
    controls, or replacement characters into extracted text. Thai combining
    marks can also be detached or duplicated. These are extraction failures,
    not spelling mistakes, and should trigger OCR or be excluded from findings.
    """
    for char in text:
        if char == "\ufffd":
            return True
        if char not in "\n\r\t" and unicodedata.category(char).startswith("C"):
            return True

    last_base = None
    marks_since_base = set()
    for char in text:
        if char in THAI_BASES:
            last_base = char
            marks_since_base.clear()
        elif char in THAI_MARKS:
            if last_base is None or char in marks_since_base:
                return True
            marks_since_base.add(char)
        elif not unicodedata.combining(char):
            last_base = None
            marks_since_base.clear()
    return False


def report_comparison_key(text: str) -> str:
    normalized = normalize_thai(text)
    return re.sub(r"[\s.,]+", "", normalized).casefold()


def is_reportable_finding(found: str, suggestion: str) -> bool:
    """Return whether a pair represents a real textual difference."""
    found = normalize_thai(found)
    suggestion = normalize_thai(suggestion)
    return (
        len(found) <= 500
        and len(suggestion) <= 500
        and not has_unicode_artifacts(found)
        and not has_unicode_artifacts(suggestion)
        and report_comparison_key(found) != report_comparison_key(suggestion)
    )


def create_highlighted_pdf(input_path: Path, findings, output_path: Path) -> tuple[int, int]:
    """Highlight exact matches and return annotation and matched-finding counts."""
    doc = fitz.open(input_path)
    count = 0
    matched = 0
    seen = set()
    try:
        for finding in findings:
            try:
                page = doc[int(finding.page) - 1]
            except (ValueError, IndexError):
                continue
            rects = page.search_for(finding.found)
            if rects:
                matched += 1
            for rect in rects:
                key = (page.number, *(round(value, 2) for value in rect))
                if key in seen:
                    continue
                seen.add(key)
                annot = page.add_highlight_annot(rect)
                annot.set_info(
                    title="Spell Check",
                    content=f"{finding.found} -> {finding.suggestion}\n{finding.reason}",
                )
                annot.update()
                count += 1
        doc.save(output_path, garbage=4, deflate=True)
        return count, matched
    finally:
        doc.close()


def highlight_run_matches(paragraph, needle: str) -> int:
    """Highlight runs whose text contains needle verbatim. Matches split across
    run boundaries (different formatting) are not detected, same limitation as
    PDF's page.search_for()."""
    count = 0
    for run in paragraph.runs:
        if needle and needle in run.text:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            count += run.text.count(needle)
    return count


def locate_docx_paragraphs(doc: Document, location: str):
    """Resolve a location label (from process_docx_unit) back to the docx
    paragraph object(s) it came from, so highlighting stays targeted."""
    match = re.match(r"^¶(\d+)$", location or "")
    if match:
        index = int(match.group(1)) - 1
        if 0 <= index < len(doc.paragraphs):
            return [doc.paragraphs[index]]
        return []
    match = re.match(r"^Table(\d+) R(\d+)C(\d+)$", location or "")
    if match:
        t_index, r_index, c_index = (int(value) - 1 for value in match.groups())
        if t_index < len(doc.tables):
            table = doc.tables[t_index]
            if r_index < len(table.rows):
                row = table.rows[r_index]
                if c_index < len(row.cells):
                    return list(row.cells[c_index].paragraphs)
    return None


def all_docx_paragraphs(doc: Document):
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def create_highlighted_docx(input_path: Path, findings, output_path: Path) -> tuple[int, int]:
    """Highlight exact matches in a DOCX and return annotation and matched-finding counts."""
    doc = Document(str(input_path))
    count = 0
    matched = 0
    for finding in findings:
        paragraphs = locate_docx_paragraphs(doc, finding.page)
        if paragraphs is None:
            paragraphs = all_docx_paragraphs(doc)
        found_here = 0
        for paragraph in paragraphs:
            if finding.found in paragraph.text:
                found_here += highlight_run_matches(paragraph, finding.found)
        if found_here:
            matched += 1
            count += found_here
    doc.save(output_path)
    return count, matched


def iter_docx_units(path: Path) -> list[tuple[str, str]]:
    """Extract (location, text) units from paragraphs and table cells, in
    document order. Nested tables and merged-cell duplicates are not
    de-duplicated (best-effort location tracking)."""
    doc = Document(str(path))
    units = []
    for index, paragraph in enumerate(doc.paragraphs, 1):
        text = paragraph.text.strip()
        if text:
            units.append((f"¶{index}", text))
    for t_index, table in enumerate(doc.tables, 1):
        for r_index, row in enumerate(table.rows, 1):
            for c_index, cell in enumerate(row.cells, 1):
                text = cell.text.strip()
                if text:
                    units.append((f"Table{t_index} R{r_index}C{c_index}", text))
    return units


def process_docx_unit(order: int, location: str, text: str, settings, fixes: list[tuple[str, str]]) -> tuple[int, str, str, list[str], str]:
    """Same normalize/chunk/dictionary/LLM steps as process_page, for a pre-extracted docx unit."""
    text = normalize_thai(text)
    chunks = chunk_text(text, MAX_CHARS)
    lines = dictionary_lines_mem(text, fixes)

    for part, chunk in enumerate(chunks, 1):
        llm = clean_report(ai_check(chunk, location, part, len(chunks)))
        if llm != "ไม่พบคำผิดชัดเจน":
            lines.append(llm)

    unit_report = clean_report("\n".join(lines))
    ocr_chunk = f"\n\n<!-- {location} -->\n\n{text.strip()}\n"

    return order, location, ocr_chunk, lines, unit_report


def split_long_block(block: str, limit: int) -> list[str]:
    parts = []
    text = block.strip()
    while len(text) > limit:
        window = text[:limit]
        cut = max(window.rfind("\n"), window.rfind(" "), window.rfind("।"), window.rfind("ฯ"), window.rfind("."), window.rfind(";"))
        if cut < limit * 0.6:
            cut = limit
        parts.append(text[:cut].strip())
        text = text[cut:].strip()
    if text:
        parts.append(text)
    return parts


def chunk_text(text: str, limit: int = MAX_CHARS) -> list[str]:
    blocks = re.split(r"\n\s*\n", text.strip())
    chunks = []
    current = ""

    for block in blocks:
        block = block.strip()
        if not block:
            continue
        candidates = split_long_block(block, limit) if len(block) > limit else [block]
        for candidate in candidates:
            if len(current) + len(candidate) + 2 <= limit:
                current = f"{current}\n\n{candidate}".strip()
            else:
                if current:
                    chunks.append(current)
                current = candidate

    if current:
        chunks.append(current)
    return chunks


def dictionary_lines_mem(text: str, fixes: list[tuple[str, str]]) -> list[str]:
    lines = []
    for wrong, correct in fixes:
        if wrong in text and not text.startswith(correct, text.find(wrong)):
            lines.append(f"- `{wrong}` -> `{correct}` : dictionary")
    return lines


def error_summary(exc: Exception) -> str:
    text = str(exc)
    if "String or binary data would be truncated" in text:
        column = re.search(r"column '([^']+)'", text)
        return f"บันทึกข้อมูลไม่สำเร็จ: ข้อมูลยาวเกินขนาดคอลัมน์ฐานข้อมูล{f' {column.group(1)}' if column else ''}"
    if isinstance(exc, RuntimeError) and "COPILOT_STUDIO_SECRET" in text:
        return "ยังไม่ได้ตั้งค่า COPILOT_STUDIO_SECRET สำหรับ AI_PROVIDER=copilot_studio"
    if isinstance(exc, urllib.error.HTTPError):
        return f"ติดต่อ API ภายนอกไม่สำเร็จ: HTTP {exc.code} {exc.reason}"
    if getattr(exc, "status_code", None) == 429 or type(exc).__name__ == "RateLimitError":
        return "OpenRouter จำกัดอัตราคำขอ (rate limit) ชั่วคราว กรุณาลองใหม่อีกครั้งภายหลัง"
    if getattr(exc, "status_code", None):
        return f"ติดต่อ API ภายนอกไม่สำเร็จ: HTTP {exc.status_code}"
    if type(exc).__name__ in {"APIConnectionError", "APITimeoutError"}:
        return f"ติดต่อ API ภายนอกไม่สำเร็จ: {type(exc).__name__}"
    if isinstance(exc, (urllib.error.URLError, TimeoutError)):
        return f"ติดต่อ API ภายนอกไม่สำเร็จ: {text}"
    return f"เกิดข้อผิดพลาดระหว่างประมวลผล: {type(exc).__name__}"


def log(db: Session, job: Job, message: str, level: str = "INFO", detail: str | None = None):
    uploader = getattr(getattr(job, "user", None), "email", None) or f"user_id={job.user_id}"
    context = f"job_id: {job.id}\nไฟล์: {job.original_filename}\nผู้อัปโหลด: {uploader}"
    logger.log(getattr(logging, level, logging.INFO), "%s\n%s", message, f"{context}\n\n{detail}" if detail else context)


def ocr_page(pdf_path: Path, page: int) -> str:
    settings = get_settings()
    client = OpenAI(
        base_url=settings.typhoon_base_url,
        api_key=settings.typhoon_ocr_api_key,
        timeout=120,
        max_retries=2,
    )
    response = client.chat.completions.create(
        model="typhoon-ocr",
        messages=prepare_ocr_messages(str(pdf_path), task_type="v1.5", page_num=page, figure_language="Thai"),
        max_tokens=16384,
        extra_body={"repetition_penalty": 1.1, "temperature": 0.1, "top_p": 0.6},
    )
    return response.choices[0].message.content


def smart_extract_page(pdf_path: Path, page_num: int) -> str:
    """Extract text from a single PDF page using pymupdf."""
    doc = fitz.open(str(pdf_path))
    try:
        page = doc[page_num - 1]  # fitz uses 0-indexed pages
        return page.get_text().strip()
    finally:
        doc.close()


def openrouter_check(text: str, page: int, part: int = 1, total: int = 1) -> str:
    settings = get_settings()
    
    # กรอง pattern ปี พ.ศ. 25\d{2} ออกก่อนส่งเพื่อประหยัด token
    clean_text = re.sub(r"\b25\d{2}\b", "", text)
    
    prompt = f"""ตรวจคำผิดภาษาไทยจาก OCR หน้า {page} (ช่วงที่ {part}/{total})

ข้อความ:
```text
{clean_text}
```"""
    client = OpenAI(
        base_url="https://openrouter.ai/api/v1",
        api_key=settings.open_router_api_key,
        timeout=120,
        max_retries=2,
    )
    response = client.chat.completions.create(
        model=settings.openrouter_model,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
        temperature=0.1,
    )
    return (response.choices[0].message.content or "ไม่พบคำผิดชัดเจน").strip()


DIRECTLINE_BASE = "https://directline.botframework.com/v3/directline"
COPILOT_STUDIO_CALL_TIMEOUT = 30
COPILOT_STUDIO_POLL_TIMEOUT_SECONDS = 20
COPILOT_STUDIO_USER_ID = "spell-check-backend"


def _directline_call(method: str, path: str, secret: str, body: dict | None = None) -> dict:
    url = f"{DIRECTLINE_BASE}{path}"
    data = json.dumps(body).encode("utf-8") if body is not None else None
    request = urllib.request.Request(
        url,
        data=data,
        method=method,
        headers={"Authorization": f"Bearer {secret}", "Content-Type": "application/json"},
    )
    with urllib.request.urlopen(request, timeout=COPILOT_STUDIO_CALL_TIMEOUT) as response:
        return json.loads(response.read().decode("utf-8"))


def copilot_studio_check(text: str, page, part: int = 1, total: int = 1) -> str:
    """Send the same spell-check prompt used for OpenRouter to a published
    Copilot Studio agent over Direct Line (same flow as the connectivity
    test), and return its raw text reply for clean_report() to parse."""
    settings = get_settings()
    secret = settings.copilot_studio_secret
    if not secret:
        raise RuntimeError("COPILOT_STUDIO_SECRET is not set")

    clean_text = re.sub(r"\b25\d{2}\b", "", text)
    prompt = f"""ตรวจคำผิดภาษาไทยจาก OCR หน้า {page} (ช่วงที่ {part}/{total})

ข้อความ:
```text
{clean_text}
```"""
    full_message = f"{SYSTEM_PROMPT}\n\n{prompt}"

    logger.info("Copilot Studio: checking page=%s part=%s/%s (%d chars)", page, part, total, len(text))
    conversation = _directline_call("POST", "/conversations", secret)
    conversation_id = conversation["conversationId"]

    _directline_call(
        "POST",
        f"/conversations/{conversation_id}/activities",
        secret,
        body={"type": "message", "from": {"id": COPILOT_STUDIO_USER_ID}, "text": full_message, "locale": "th-TH"},
    )

    watermark = None
    deadline = time.time() + COPILOT_STUDIO_POLL_TIMEOUT_SECONDS
    while time.time() < deadline:
        time.sleep(1.5)
        path = f"/conversations/{conversation_id}/activities"
        if watermark:
            path += f"?watermark={watermark}"
        payload = _directline_call("GET", path, secret)
        watermark = payload.get("watermark") or watermark
        for activity in payload.get("activities", []):
            if activity.get("type") == "message" and activity.get("from", {}).get("id") != COPILOT_STUDIO_USER_ID and activity.get("text"):
                logger.info("Copilot Studio: got reply for page=%s part=%s/%s", page, part, total)
                return activity["text"].strip()

    raise TimeoutError(f"Copilot Studio ไม่ตอบกลับภายใน {COPILOT_STUDIO_POLL_TIMEOUT_SECONDS} วินาที")


def ai_check(text: str, page, part: int = 1, total: int = 1) -> str:
    """Provider dispatcher: AI_PROVIDER=copilot_studio | openrouter (default).
    OpenRouter stays the default/fallback provider; a Copilot Studio failure
    fails this unit the same way an OpenRouter failure already does (no
    silent per-call failover, so provider comparisons stay honest). Copilot
    Studio occasionally answers "ไม่พบคำผิดชัดเจน" even for text with real
    typos (non-deterministic agent response, confirmed by repeated identical
    calls) -- a single zero-finding response is retried once with a fresh
    Direct Line conversation before being accepted."""
    settings = get_settings()
    provider = (settings.ai_provider or "openrouter").strip().lower()
    if provider == "copilot_studio":
        try:
            first = copilot_studio_check(text, page, part, total)
            if clean_report(first) != "ไม่พบคำผิดชัดเจน":
                return first
            logger.info(
                "Copilot Studio: zero findings on first attempt (page=%s part=%s/%s), retrying once with a new conversation",
                page, part, total,
            )
            second = copilot_studio_check(text, page, part, total)
            if clean_report(second) != "ไม่พบคำผิดชัดเจน":
                logger.info("Copilot Studio: retry found results (page=%s part=%s/%s)", page, part, total)
                return second
            logger.info("Copilot Studio: retry also returned zero findings (page=%s part=%s/%s), keeping zero", page, part, total)
            return second
        except Exception:
            logger.exception("Copilot Studio check failed (page=%s part=%s/%s)", page, part, total)
            raise
    return openrouter_check(text, page, part, total)


def clean_report(text: str) -> str:
    kept = []
    for line in text.splitlines():
        m = re.match(r"\s*-\s*`([^`]+)`\s*->\s*`([^`]+)`\s*(?::\s*(.+))?", line)
        if m:
            found = normalize_thai(m.group(1))
            suggestion = normalize_thai(m.group(2))
            reason = (m.group(3) or "")[:1000]

            # Broken PDF glyph mappings are extraction failures, not spelling
            # errors. Filtering by Unicode shape handles every affected word
            # instead of growing a hard-coded dictionary.
            if not is_reportable_finding(found, suggestion):
                continue
            
            # กรองข้อเสนอแนะที่เป็นเพียงปัญหาจากการอ่านของ OCR (เมื่อผู้ใช้ระบุในเหตุผล)
            if "ocr" in reason.lower() or "อ่านผิด" in reason:
                continue
                
            line = f"- `{found}` -> `{suggestion}`"
            if reason:
                line += f" : {reason}"
        kept.append(line)
    return "\n".join(kept).strip() or "ไม่พบคำผิดชัดเจน"


def text_check_findings(text: str, fixes: list[tuple[str, str]]) -> list[dict[str, str]]:
    text = normalize_thai(text)
    lines = dictionary_lines_mem(text, fixes)
    llm = clean_report(ai_check(text, 1))
    if llm != "ไม่พบคำผิดชัดเจน":
        lines.append(llm)
    findings = []
    for index, line in enumerate(clean_report("\n".join(lines)).splitlines(), 1):
        match = re.match(r"- `([^`]+)` -> `([^`]+)`(?: : (.+))?", line)
        if match:
            findings.append({
                "id": index,
                "page": "ข้อความ",
                "found": match.group(1),
                "suggestion": match.group(2),
                "reason": match.group(3) or "ตรวจพบโดยระบบ",
            })
    return findings


def process_page(pdf_path: Path, page: int, pages: int, settings, fixes: list[tuple[str, str]]) -> tuple[int, str, list[str], str]:
    # 1. Extract text (can be slow if OCR, fast if pymupdf)
    text = smart_extract_page(pdf_path, page) if settings.use_pymupdf else ocr_page(pdf_path, page)
    # 2. Normalize
    text = normalize_thai(text)
    
    # 3. Chunk text if it exceeds MAX_CHARS
    chunks = chunk_text(text, MAX_CHARS)
    
    # 4. Dictionary lookup (in-memory, fast)
    lines = dictionary_lines_mem(text, fixes)
    
    # 5. LLM checks for all chunks
    for part, chunk in enumerate(chunks, 1):
        llm = clean_report(ai_check(chunk, page, part, len(chunks)))
        if llm != "ไม่พบคำผิดชัดเจน":
            lines.append(llm)
            
    page_report = clean_report("\n".join(lines))
    ocr_chunk = f"\n\n<!-- page {page} -->\n\n{text.strip()}\n"
    
    return page, page, ocr_chunk, lines, page_report


def report_to_excel(report: str, path: Path):
    wb = Workbook()
    ws = wb.active
    ws.title = "spellcheck"
    ws.append(["page", "found", "suggestion", "reason"])
    page = ""
    for line in report.splitlines():
        if line.startswith("## Page"):
            page = line.replace("## Page", "").strip()
        m = re.match(r"- `([^`]+)` -> `([^`]+)` : (.+)", line)
        if m:
            ws.append([page, m.group(1), m.group(2), m.group(3)])
    wb.save(path)


def save_findings(db: Session, job: Job, report: str):
    db.query(SpellcheckFinding).filter(SpellcheckFinding.job_id == job.id).update(
        {SpellcheckFinding.is_active: False, SpellcheckFinding.updated_by: job.user_id},
        synchronize_session=False,
    )
    page = ""
    for line in report.splitlines():
        if line.startswith("## Page"):
            page = line.replace("## Page", "").strip()
        m = re.match(r"- `([^`]+)` -> `([^`]+)` : (.+)", line)
        if m:
            db.add(SpellcheckFinding(job_id=job.id, page=page, found=m.group(1), suggestion=m.group(2), reason=m.group(3), created_by=job.user_id, updated_by=job.user_id))
    db.commit()


def run_job(db: Session, job: Job):
    started = time.time()
    job.status = JobStatus.PROCESSING.value
    db.commit()
    log(db, job, "เริ่มประมวลผลเอกสาร")
    try:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / job.original_filename
            download_file(job.original_key, source)
            is_docx = source.suffix.lower() == ".docx"

            # Load fixes in-memory
            fixes = [(term.wrong, term.correct) for term in db.query(DictionaryTerm).filter(DictionaryTerm.is_active == True).order_by(DictionaryTerm.wrong).all()]

            settings = get_settings()

            # Submit concurrent unit processing (concurrency capped to avoid OpenRouter rate limit)
            futures = {}
            executor = ThreadPoolExecutor(max_workers=PAGE_CONCURRENCY)
            try:
                if is_docx:
                    units = iter_docx_units(source)
                    total_units = len(units)
                    job.pages = total_units
                    db.commit()
                    for order, (location, text) in enumerate(units, 1):
                        futures[executor.submit(process_docx_unit, order, location, text, settings, fixes)] = location
                else:
                    pages = len(PdfReader(str(source)).pages)
                    total_units = pages
                    job.pages = pages
                    db.commit()
                    for page in range(1, pages + 1):
                        futures[executor.submit(process_page, source, page, pages, settings, fixes)] = page

                results = []
                try:
                    for future in as_completed(futures, timeout=JOB_TIMEOUT_SECONDS):
                        label = futures[future]
                        try:
                            res = future.result()
                            results.append(res)
                            if is_docx:
                                log(db, job, f"Finished processing {label} ({len(results)}/{total_units})")
                            else:
                                log(db, job, f"Finished processing page {label}/{total_units}")
                        except Exception as exc:
                            detail = f"{error_summary(exc)}\n\n{traceback.format_exc()}"
                            if is_docx:
                                log(db, job, f"ประมวลผล {label} ไม่สำเร็จ: {error_summary(exc)}", "ERROR", detail)
                                log(db, job, f"Error processing {label}: {exc}", "ERROR")
                            else:
                                log(db, job, f"ประมวลผลหน้า {label} ไม่สำเร็จ: {error_summary(exc)}", "ERROR", detail)
                                log(db, job, f"Error processing page {label}: {exc}", "ERROR")
                            raise exc
                except TimeoutError:
                    log(db, job, f"หมดเวลาประมวลผล: เกิน {JOB_TIMEOUT_SECONDS} วินาที (อาจติด rate limit จาก OpenRouter)", "ERROR")
                    raise TimeoutError(f"งานใช้เวลานานเกิน {JOB_TIMEOUT_SECONDS} วินาที ระบบยกเลิกงานเพื่อไม่ให้ค้าง")
            finally:
                executor.shutdown(wait=False, cancel_futures=True)

            # Sort results by original order (page number, or docx unit order)
            results.sort(key=lambda x: x[0])

            ocr_chunks = []
            report = [f"# Spellcheck Report: {job.original_filename}", ""]

            for _, label, ocr_chunk, lines, unit_report in results:
                ocr_chunks.append(ocr_chunk)
                report.append(f"## Page {label}\n\n{unit_report}\n")

            ocr_path = root / f"{Path(job.original_filename).stem}.ocr.md"
            report_path = root / f"{Path(job.original_filename).stem}.spellcheck.md"
            excel_path = root / f"{Path(job.original_filename).stem}.spellcheck.xlsx"
            ocr_path.write_text("".join(ocr_chunks).strip() + "\n", encoding="utf-8")
            report_text = "\n".join(report).strip() + "\n"
            report_path.write_text(report_text, encoding="utf-8")
            report_to_excel(report_text, excel_path)
            save_findings(db, job, report_text)

            base = f"jobs/{job.id}"
            job.ocr_key = f"{base}/{ocr_path.name}"
            job.report_key = f"{base}/{report_path.name}"
            job.excel_key = f"{base}/{excel_path.name}"
            upload_file(ocr_path, job.ocr_key, "text/markdown")
            upload_file(report_path, job.report_key, "text/markdown")
            upload_file(excel_path, job.excel_key, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            job.status = JobStatus.DONE.value
            job.elapsed_seconds = int(time.time() - started)
            job.updated_by = job.user_id
            db.commit()
            unit_word = "หน่วย" if is_docx else "หน้า"
            log(db, job, f"ประมวลผลเอกสารสำเร็จ {total_units} {unit_word} ใช้เวลา {job.elapsed_seconds} วินาที")
            finding_count = db.query(SpellcheckFinding).filter(SpellcheckFinding.job_id == job.id, SpellcheckFinding.is_active == True).count()
            record_audit(
                db, AuditEvent.JOB_DONE, actor_user_id=job.user_id, job_id=job.id,
                detail=f"findings={finding_count};elapsed={job.elapsed_seconds}s",
            )
    except Exception as exc:
        detail = f"{error_summary(exc)}\n\n{traceback.format_exc()}"
        db.rollback()
        job.status = JobStatus.FAILED.value
        job.error_text = str(exc)[:1000]
        job.elapsed_seconds = int(time.time() - started)
        job.updated_by = job.user_id
        db.commit()
        log(db, job, job.error_text, "ERROR", detail)
        record_audit(
            db, AuditEvent.JOB_FAILED, actor_user_id=job.user_id, job_id=job.id,
            detail=error_summary(exc)[:500],
        )
